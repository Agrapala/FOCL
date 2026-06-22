"""
generate_mechanism_doc.py
Produces: results/FLoBC_PoCL_Mechanism.docx
Full system description for block diagram creation — no code changes.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_PATH = os.path.join(os.path.dirname(__file__), "results", "FLoBC_PoCL_Mechanism.docx")
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)


# ── Helper functions ──────────────────────────────────────────────────────────

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x37, 0x64, 0x92)
    return p

def body(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.3 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    return p

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def label(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(6)
    return p

def note(text):
    p = doc.add_paragraph()
    run = p.add_run("NOTE: " + text)
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "BDD7EE")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(cell_text)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

def divider():
    doc.add_paragraph("─" * 90).runs[0].font.size = Pt(7)


# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("FLoBC-PoCL System Mechanism")
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run(
    "Blockchain-Based Federated Learning with Proof-of-Collaborative-Learning\n"
    "for Multi-Hospital Pneumonia Diagnosis\n\n"
    "Full Mechanism Document for Block Diagram Construction"
)
sub_run.font.size = Pt(13)
sub_run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — SYSTEM OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════

h1("1. System Overview")

body(
    "FLoBC-PoCL is a privacy-preserving, decentralized federated learning (FL) "
    "framework built on a blockchain infrastructure. Four Sri Lankan hospitals "
    "(nodes) collaboratively train a shared pneumonia detection model without "
    "ever sharing raw patient X-ray images. All coordination, audit trail, and "
    "model aggregation decisions are recorded on an immutable, cryptographically "
    "linked blockchain."
)

body(
    "The system extends the original FLoBC framework (Abuzied et al., Cluster "
    "Computing 2024) by replacing naive Federated Averaging with a "
    "Proof-of-Collaborative-Learning (PoCL) consensus mechanism that selects "
    "only the top-performing miners per round for aggregation, rewarding "
    "contribution quality and submission timeliness."
)

doc.add_paragraph()
label("Four Participating Hospitals (Nodes)")
add_table(
    ["Node ID", "Hospital Name",     "City",   "Blockchain Address (prefix)"],
    [
        ["A",   "Hospital_Galle",   "Galle",   "0xAAA...0001"],
        ["B",   "Hospital_Colombo", "Colombo", "0xBBB...0002"],
        ["C",   "Hospital_Kandy",   "Kandy",   "0xCCC...0003"],
        ["D",   "Hospital_Jaffna",  "Jaffna",  "0xDDD...0004"],
    ],
    col_widths=[0.8, 1.8, 1.2, 2.5],
)

label("Key Privacy Guarantee")
body(
    "Each hospital trains on its own PRIVATE chest X-ray images stored locally. "
    "Only the SHA-256 hash of the trained model weights is broadcast on-chain — "
    "the actual weight values are shared in-memory between components on the "
    "same machine but never transmitted as raw patient data."
)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — SYSTEM COMPONENTS
# ═══════════════════════════════════════════════════════════════════════════════

h1("2. System Components")

h2("2.1 Data Layer — Pneumonia Loader (pneumonia_loader.py)")

body("Responsible for loading real chest X-ray images and constructing data splits.")

label("Image Preprocessing Pipeline")
numbered("Load JPEG/PNG images from per-hospital folders (Node_A, Node_B, Node_C, Node_D).")
numbered("Convert each image to grayscale (L channel) using Pillow.")
numbered("Resize to 64 × 64 pixels → 4,096 pixel features per image.")
numbered("Normalise pixel values to [0, 1] (divide by 255).")
numbered("Flatten to a 1-D float32 vector of length 4,096.")
numbered("Label: 0 = NORMAL, 1 = PNEUMONIA.")

label("Data Split Strategy (build_splits)")
body(
    "After loading, each hospital's data is independently split. "
    "No hospital's training data is ever shared with another hospital."
)
add_table(
    ["Split", "Ratio", "Purpose", "Shared?"],
    [
        ["Train",              "75%",  "Local SGD training at each hospital",                "NO — stays private"],
        ["Validation (local)", "15%",  "Per-trainer early stopping during local training",    "NO"],
        ["Validation (pooled)","15%*", "BC validators judge incoming model updates",          "YES — pooled across 4 nodes"],
        ["Test (per-node)",    "10%",  "Measure per-hospital accuracy before/after FL",       "NO"],
        ["Test (pooled)",      "10%*", "Global model evaluation after each FL round",         "YES — pooled across 4 nodes"],
    ],
    col_widths=[1.6, 0.7, 2.9, 1.7],
)
note("* Pooled = concatenation of each hospital's corresponding split into one shared array.")

h2("2.2 Model Layer — PneumoniaModel (flobc_pneumonia_engine.py)")

body(
    "A two-layer Multilayer Perceptron (MLP) implemented in pure NumPy. "
    "No deep learning framework dependency."
)

label("Architecture")
add_table(
    ["Layer", "Type",        "Input Dim", "Output Dim", "Activation"],
    [
        ["1 (Hidden)", "Fully Connected", "4,096",  "256",   "ReLU"],
        ["2 (Output)", "Fully Connected", "256",    "2",     "Softmax"],
    ],
    col_widths=[1.2, 1.8, 1.2, 1.2, 1.5],
)

label("Weight Initialisation")
body(
    "He initialisation: W ~ N(0, sqrt(2/fan_in)) for numerical stability "
    "with ReLU activations. Biases initialised to zero."
)

label("Training — SGD with Mini-batches")
bullet("Mini-batch SGD with configurable batch size (default: 32).")
bullet("Learning rate: configurable per run (default: 0.008).")
bullet("Early stopping: patience = 4 epochs; best weights are saved.")
bullet("freeze_features=True mode: only W2, b2 (output head) updated — "
       "W1, b1 (feature extractor) frozen. Used during personalisation fine-tuning.")

label("Flat Weight Representation")
body(
    "All four weight matrices (W1, b1, W2, b2) are concatenated into a single "
    "1-D NumPy array (flat vector) for network transmission, hashing, and "
    "FedAvg aggregation. Total parameter count: "
    "4096×256 + 256 + 256×2 + 2 = 1,049,602 parameters."
)

h2("2.3 Hospital Trainer Node — HospitalTrainer")

body(
    "One HospitalTrainer instance runs at each hospital. It owns a local copy "
    "of the MLP, a cryptographic Wallet, and trains on private local data."
)

label("Per-Round Local Training Sequence")
numbered("Pull global model: compare global vs local accuracy on local val set; "
         "accept global only if within 2% of local (pull_global method).")
numbered("Run local_epochs SGD epochs with mini-batch shuffling.")
numbered("Track best validation accuracy; save best weights; apply early stopping.")
numbered("If noise_std > 0 (Byzantine simulation): add Gaussian noise to submitted weights.")
numbered("Compute SHA-256 hash of submitted weights.")
numbered("Create and sign MODEL_UPDATE transaction via RSA-2048 wallet.")
numbered("Return: (weights, weights_hash, signed_tx, train_acc, val_acc).")

note(
    "Only the hash goes on-chain. Weight values are passed in-memory "
    "within the same Python process. In a real distributed deployment, "
    "weights would be transmitted over a secure channel."
)

h2("2.4 Blockchain Validator Node — BCValidator")

body(
    "Three independent BCValidator instances judge incoming model weight updates. "
    "Each validator has its own slice of the pooled validation set and its own Wallet."
)

label("Validation Rule")
body(
    "For each candidate model update, the validator loads the weights into a "
    "temporary model copy and measures accuracy on its local validation slice. "
    "Decision threshold: max(0.45, baseline_accuracy - 0.08). "
    "The update is ACCEPTED if its accuracy meets this threshold; REJECTED otherwise."
)
bullet("ACCEPTED → update enters the candidate pool for aggregation.")
bullet("REJECTED → update is excluded from FedAvg this round.")
bullet("Threshold is intentionally forgiving in early rounds (when accuracy "
       "is still below 0.53) to avoid blocking valid early updates.")

label("Block-Level Consensus (pBFT / Proof-of-Stake)")
body(
    "After all validators cast their per-update votes, a separate PoS consensus "
    "round determines whether the entire block (containing all round transactions) "
    "is committed to the chain. A block is finalized only if validators holding "
    "> 2/3 of total stake vote YES."
)

h2("2.5 Trust / Reputation Service — TrustService")

body(
    "Maintains a normalized reputation score for each hospital trainer. "
    "Scores are used as weights in reputation-weighted FedAvg aggregation."
)

label("Score Update Rules (per round, per trainer)")
add_table(
    ["Condition",                          "Action",                       "Delta"],
    [
        ["delta_acc > 0  (improvement)",   "REWARD  score += 0.08",        "+0.08"],
        ["delta_acc < -0.01  (regression)","PENALTY  score -= 0.10",       "-0.10"],
        ["-0.01 <= delta_acc <= 0",        "Neutral — no change",          "0"],
    ],
    col_widths=[2.8, 2.2, 0.8],
)
body("delta_acc = candidate_accuracy - baseline_global_accuracy on pooled test set.")
bullet("Minimum score floor: 0.05 (no healthy hospital collapses to zero).")
bullet("After each update, all scores are renormalized to sum to 1.0.")
bullet("Byzantine node (simulated with large noise) will persistently receive "
       "PENALTY, collapsing its weight toward the 0.05 floor.")

h2("2.6 Cryptographic Layer (blockchain/crypto.py)")

label("SHA-256 Hashing")
bullet("sha256(str) — UTF-8 string hashing for transaction and block hashes.")
bullet("sha256_bytes(bytes) — raw bytes hashing for model weight fingerprinting.")
bullet("double_sha256 — Bitcoin-style double hash (available but not used in main path).")

label("Merkle Tree")
body(
    "Each block computes a Merkle root over all its transaction hashes. "
    "Any tampering with a single transaction changes the Merkle root, "
    "which then changes the block hash, breaking the chain link. "
    "Odd-length lists are padded by duplicating the last hash (Bitcoin convention)."
)

label("Wallet (RSA-2048 Digital Signatures)")
bullet("Each node (trainer + validator) generates an RSA-2048 key pair at startup.")
bullet("Private key signs transaction hashes (PKCS1v15 + SHA-256).")
bullet("Public key hash (SHA-256 of PEM-encoded public key) = node address.")
bullet("Fallback: HMAC-SHA-256 pseudo-signing if the cryptography library is unavailable.")

h2("2.7 Transaction Layer (blockchain/transaction.py)")

body("Every event in the FL process is atomically recorded as a signed Transaction.")

add_table(
    ["TX Type",             "Who Creates",          "Key Payload Fields",                               "Phase"],
    [
        ["MODEL_UPDATE",       "Hospital Trainer",  "trainer_id, round_num, weights_hash, noise_level", "A"],
        ["PREDICTION_PROPOSAL","Hospital Trainer",  "trainer_id, round_num, predictions_hash, elapsed_s","B (PoCL)"],
        ["VOTE",               "BC Validator",      "voter_id, trainer_id, accuracy_score, timeliness_score, vote_score", "C (PoCL)"],
        ["WINNER_SELECTION",   "FL Framework",      "round_num, winners[], vote_scores{}",              "D (PoCL)"],
        ["TRUST_UPDATE",       "FL Framework",      "trainer_id, old_score, new_score, round_num",      "E"],
        ["REWARD",             "Hospital Trainer",  "trainer_id, round_num, contribution (R_i)",        "E (PoCL)"],
        ["GLOBAL_MODEL",       "FL Framework",      "round_num, weights_hash, accuracy, accepted_trainers[], trust_scores{}", "F"],
        ["VALIDATION",         "BC Validator",      "validator_id, trainer_id, accepted, score",        "C (standard FL)"],
    ],
    col_widths=[1.6, 1.5, 3.2, 0.8],
)

label("Transaction Lifecycle")
numbered("Created with payload data by the originating component.")
numbered("Timestamp set to current Unix epoch.")
numbered("Canonical JSON serialized (sorted keys, no whitespace) for deterministic hashing.")
numbered("SHA-256 hash computed and stored as tx_hash.")
numbered("RSA-2048 signature computed over tx_hash by creator's private key.")
numbered("Transaction broadcast to validators for inclusion in next block.")

h2("2.8 Block and Chain Layer (blockchain/chain.py)")

label("Block Structure")
add_table(
    ["Field",          "Type",   "Description"],
    [
        ["index",          "int",    "Position in chain (0 = genesis)"],
        ["previous_hash",  "str",    "SHA-256 hash of parent block — tamper-evident link"],
        ["transactions",   "list",   "All signed Transaction objects for this round"],
        ["merkle_root",    "str",    "Merkle root of all transaction hashes"],
        ["validator",      "str",    "Address of the PoS-selected proposer"],
        ["stake_votes",    "dict",   "{validator_address: stake_weight} — PoS YES votes"],
        ["timestamp",      "float",  "Unix epoch when block was proposed"],
        ["nonce",          "int",    "Light PoW fallback (unused in PoS-only path)"],
        ["block_hash",     "str",    "SHA-256 of all above fields (header string)"],
    ],
    col_widths=[1.5, 0.8, 4.3],
)

label("Proof-of-Stake (PoS) Consensus — Block Level")
numbered("Validator with highest stake is selected as block proposer.")
numbered("Candidate block is built with all round transactions and merkle_root computed.")
numbered("All validators independently verify block structural integrity.")
numbered("Each validator that agrees casts a YES vote with its stake weight.")
numbered("If YES stake > 2/3 of total stake: block is finalized (pBFT threshold).")
numbered("Block committed to chain; pending transaction pool cleared.")

label("Chain Integrity Invariants")
bullet("block[i].previous_hash == block[i-1].block_hash for all i >= 1.")
bullet("block[i].block_hash == SHA-256(header_string) for all blocks.")
bullet("block[i].merkle_root == MerkleTree(tx_hashes).root for all blocks.")
bullet("All transactions pass tx_hash == SHA-256(canonical_json) check.")
bullet("stake_votes represent > 2/3 of total stake for all non-genesis blocks.")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — STANDARD FL ROUND (BASELINE)
# ═══════════════════════════════════════════════════════════════════════════════

h1("3. Standard FL Round Mechanism (Baseline FLoBC)")

body(
    "The standard FL training loop (train() method) follows a sequential "
    "7-phase protocol per round. This is the baseline before PoCL enhancement."
)

h2("Phase A — Local Training at Each Hospital")
body("Each of the 4 hospitals independently trains its local MLP on private data.")
bullet("Hospital pulls the latest global model (with pull_global acceptance check).")
bullet("Runs local_epochs SGD epochs with mini-batch shuffling and early stopping.")
bullet("Best weights (by local validation accuracy) are saved.")
bullet("Submits: weights vector (in-memory), SHA-256 hash, signed MODEL_UPDATE tx.")

h2("Phase B — Synchronisation Filter")
body("Determines which hospital updates proceed to validation this round.")
add_table(
    ["Scheme", "Behaviour",                                            "Parameter"],
    [
        ["BSP (Bulk Synchronous Parallel)",    "ALL 4 hospitals must submit before proceeding (default)", "—"],
        ["SSP (Stale Synchronous Parallel)",   "Top (1 - slack_ratio) fraction of hospitals proceed",    "ssp_slack_ratio = 0.2"],
        ["BAP (Barrierless Async Parallel)",   "Any majority_ratio fraction of hospitals can proceed",   "bap_majority_ratio = 1.0"],
    ],
    col_widths=[2.2, 3.2, 1.5],
)

h2("Phase C — Blockchain Validation (>2/3 Consensus)")
body("For each candidate hospital update:")
bullet("All 3 BCValidators independently evaluate candidate weights on their local val slices.")
bullet("Decision: ACCEPT if accuracy >= max(0.45, baseline - 0.08); REJECT otherwise.")
bullet("Consensus: update is accepted if more than 2/3 of validators say YES.")
bullet("Each validator creates a signed VALIDATION transaction on-chain.")

h2("Phase D — Trust Score Update")
body("For each hospital (regardless of consensus outcome):")
bullet("Compute delta_acc = (candidate accuracy on pooled test) - (current global accuracy).")
bullet("REWARD if delta > 0; PENALTY if delta < -0.01; neutral otherwise.")
bullet("Scores renormalized to sum to 1.0.")
bullet("TRUST_UPDATE transaction created and added to block.")

h2("Phase E — Reputation-Weighted Federated Averaging (FedAvg)")
body("Only consensus-accepted updates enter aggregation:")
body(
    "new_global = sum(trust_weight_i * weights_i) for all accepted hospitals i\n"
    "where trust_weight_i = TrustService.scores[i] / sum(accepted trust scores)"
)
bullet("The new global model replaces the old one.")
bullet("All hospitals and validators receive the updated global via push_global / sync_global.")
bullet("Non-accepted hospitals keep their local model for next round.")

h2("Phase F — Global Model Evaluation")
bullet("Global model accuracy measured on the pooled test set (all 4 hospitals' test slices).")
bullet("Accuracy log updated; trust score log updated.")
bullet("GLOBAL_MODEL transaction created: records weights hash, accuracy, accepted trainers, trust scores.")

h2("Phase G+H — Block Sealing (PoS/pBFT)")
bullet("All phase A–F transactions collected into block_txs list.")
bullet("PoS proposer selected (highest-stake validator).")
bullet("Block proposed, all validators vote, >2/3 stake = YES → block committed.")
bullet("51 total blocks at end of 50-round run (1 genesis + 50 FL-round blocks).")

label("Transactions per Standard FL Block (29 per round)")
add_table(
    ["TX Type",       "Count/Round", "Total (50 rounds)"],
    [
        ["MODEL_UPDATE",   "4",  "200"],
        ["VALIDATION",     "12 (4 trainers × 3 validators)", "0*"],
        ["TRUST_UPDATE",   "4",  "200"],
        ["GLOBAL_MODEL",   "1",  "50"],
    ],
    col_widths=[2.0, 2.5, 2.0],
)
note(
    "* VALIDATION tx count is 0 in PoCL mode because the VALIDATION phase is "
    "replaced by the PREDICTION_PROPOSAL + VOTE + WINNER_SELECTION protocol. "
    "In PoCL runs: MODEL_UPDATE=200, PREDICTION_PROPOSAL=200, VOTE=600, "
    "WINNER_SELECTION=50, TRUST_UPDATE=200, REWARD=150, GLOBAL_MODEL=50."
)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — PoCL CONSENSUS MECHANISM (ENHANCED)
# ═══════════════════════════════════════════════════════════════════════════════

h1("4. PoCL (Proof-of-Collaborative-Learning) Mechanism")

body(
    "PoCL replaces the standard >2/3 validator acceptance vote with a "
    "six-phase competitive consensus. Instead of aggregating all accepted "
    "hospitals, only the top-K performers (by vote score) contribute to the "
    "global model each round. This incentivises quality and timeliness, and "
    "implements the reward distribution described in FLoBC-PoCL."
)

label("PoCL Parameters")
add_table(
    ["Parameter",          "Default", "Meaning"],
    [
        ["n_rounds",           "50",    "Total FL rounds"],
        ["k_winners",          "3",     "Top-K miners selected per round"],
        ["eval_batch_size",    "64",    "Size of shared evaluation batch (sampled from pooled val set)"],
        ["timeliness_weight",  "0.3",   "Weight of timeliness in combined vote score (0.7 = accuracy weight)"],
    ],
    col_widths=[1.8, 0.8, 4.0],
)

h2("Phase A — Model Proposal")
body("Identical to standard FL Phase A (local training). Each hospital:")
bullet("Runs local SGD with early stopping.")
bullet("Records submission time (submit_times[tid] = elapsed seconds).")
bullet("Creates MODEL_UPDATE transaction on-chain.")
bullet("Returns model weights in-memory for subsequent PoCL phases.")

h2("Phase B — Prediction Proposal")
body(
    "A shared public evaluation batch is sampled from the pooled validation pool "
    "(64 samples chosen randomly each round). Every hospital runs inference with "
    "its freshly-trained model on this batch."
)
numbered("Sample 64 (= eval_batch_size) indices from pooled X_val / y_val.")
numbered("For each hospital: load its weights into a temp model; run forward pass.")
numbered("Compute SHA-256 hash of the prediction array (int32 bytes).")
numbered("Create PREDICTION_PROPOSAL transaction: predictions_hash + submit_elapsed_s.")
note(
    "The actual prediction array is not stored on-chain — only its hash. "
    "Validators re-compute accuracy independently using the labels they hold."
)

h2("Phase C — Vote Proposal")
body(
    "Each of the 3 BC validators independently scores all 4 hospitals' predictions."
)
label("Combined Vote Score Formula")
body("vote_score = (1 - timeliness_weight) * accuracy_score + timeliness_weight * timeliness_score")
body("Where:")
bullet("accuracy_score = fraction of the 64 evaluation samples correctly predicted.")
bullet("timeliness_score = 1.0 - (this_hospital_elapsed / max_elapsed_across_all_hospitals).")
bullet("timeliness_weight = 0.3 (default): accuracy has 70% weight, timeliness 30%.")
body("This produces 4 hospitals × 3 validators = 12 VOTE transactions per round.")

label("Vote Score Averaging")
body("avg_vote[tid] = mean of 3 validators' vote scores for hospital tid.")

h2("Phase D — Winner Selection")
numbered("Sort hospitals by avg_vote descending.")
numbered("Select top-K = top 3 hospitals as winners.")
numbered("Hash-verify each winner's submitted weights (integrity check).")
numbered("Record WINNER_SELECTION transaction: winners list + all vote scores.")
body(
    "Only winning hospitals' weights enter the FedAvg aggregation this round. "
    "Non-winners are not penalized beyond normal trust accounting — they "
    "simply do not receive a reward and their weights are not aggregated."
)

h2("Phase E — Reward Distribution and Trust Update")
label("Contribution Score R_i (FLoBC-PoCL Reward Formula)")
body(
    "For each winning hospital i, the contribution score R_i measures how "
    "much its local training diverged from the current global model:"
)
body(
    "R_i = (1/L) * sum_l [ (1/N_l) * sum_n | W_n^l(local) - W_n^l(global) | ]"
)
bullet("L = number of weight layers (4 layers: W1, b1, W2, b2).")
bullet("N_l = number of parameters in layer l.")
bullet("The mean absolute difference per parameter, averaged across layers.")
bullet("Higher R_i = more divergence from global model = larger unique contribution.")
body("Each winner receives a signed REWARD transaction on-chain recording its R_i score.")

label("Trust Update — All Hospitals")
body(
    "Trust scores are updated for ALL hospitals (winners and non-winners) based "
    "on delta_acc = candidate accuracy on pooled test - current global accuracy. "
    "Non-winners receive no REWARD tx but are still trust-scored."
)

h2("Phase F — Federated Averaging (Winners Only) + Block Creation")
numbered("FedAvg applied only to the top-K winning hospitals.")
numbered("Aggregation: new_global = sum(trust_weight_i * weights_i) for winning i only.")
numbered("Global model updated; all hospitals and validators synced via push_global.")
numbered("GLOBAL_MODEL transaction records the new weights hash, accuracy, accepted_trainers=winners.")
numbered("All phase A–F transactions (29 per round) sealed into one PoS/pBFT block.")

label("PoCL Transaction Count per Round (29 transactions)")
add_table(
    ["TX Type",             "Count", "Detail"],
    [
        ["MODEL_UPDATE",       "4",  "One per hospital trainer"],
        ["PREDICTION_PROPOSAL","4",  "One per hospital trainer"],
        ["VOTE",               "12", "3 validators × 4 hospitals"],
        ["WINNER_SELECTION",   "1",  "Records top-3 winners"],
        ["TRUST_UPDATE",       "4",  "One per hospital (all, not just winners)"],
        ["REWARD",             "3",  "One per winner (k=3)"],
        ["GLOBAL_MODEL",       "1",  "New aggregated global model hash + accuracy"],
        ["TOTAL",              "29", "Per round × 50 rounds = 1,450 total transactions"],
    ],
    col_widths=[1.9, 0.8, 4.0],
)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — PULL_GLOBAL SELECTIVE ACCEPTANCE
# ═══════════════════════════════════════════════════════════════════════════════

h1("5. Selective Global Model Acceptance (pull_global)")

body(
    "At the start of each round, before local training begins, each hospital "
    "decides whether to adopt the incoming federated global model or retain "
    "its own current model."
)

label("Decision Rule")
body(
    "If global_model.accuracy(local_val) >= local_model.accuracy(local_val) - 0.02:\n"
    "    Accept global (replace local with global)\n"
    "Else:\n"
    "    Keep local model"
)

body(
    "The 2% tolerance prevents high-baseline hospitals (Galle 91%, Kandy 92%) "
    "from being overwritten by a cross-hospital global model that doesn't capture "
    "their specific data distribution, while still allowing genuine improvements "
    "from federation to propagate."
)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — PERSONALIZATION (FINE-TUNING AFTER FL)
# ═══════════════════════════════════════════════════════════════════════════════

h1("6. Post-FL Personalisation (Frozen-Feature Fine-Tuning)")

body(
    "After federated training, each hospital optionally fine-tunes the global "
    "model on its own local data to adapt to its patient population. This "
    "measures whether federation provided a better starting point than training "
    "from scratch."
)

label("Personalisation Protocol (personalize_and_evaluate in evaluate_objectives.py)")
numbered("Start from the FL-trained global model weights.")
numbered("Fine-tune for ft_epochs = 40 epochs on each hospital's local training data.")
numbered("Feature extractor (W1, b1) is FROZEN — only output head (W2, b2) is updated.")
numbered("Frozen-feature fine-tuning prevents catastrophic forgetting of FL-learned "
         "representations while adapting the classifier head to local label distribution.")
numbered("Evaluate on the hospital's own held-out per_node_test set.")
numbered("Compare against solo-trained baseline (same architecture, no federation).")

label("Why Freeze the Feature Extractor?")
bullet("W1, b1 encode general pneumonia features learned from all 4 hospitals' data.")
bullet("Updating W1 on a small local dataset risks overfitting and erasing cross-hospital knowledge.")
bullet("Only the 256×2 + 2 = 514 output-head parameters are re-tuned (low variance).")
bullet("Equivalent to FedRep / FedPer style personalisation in the FL literature.")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — BYZANTINE FAULT TOLERANCE
# ═══════════════════════════════════════════════════════════════════════════════

h1("7. Byzantine Fault Tolerance")

body(
    "The system is designed to detect and isolate malicious or faulty nodes "
    "without requiring a trusted central authority."
)

label("Noise Injection (Byzantine Simulation)")
body(
    "A Byzantine hospital is simulated by adding Gaussian noise (std = 0.5) "
    "to the submitted weight vector: submitted_weights += N(0, noise_std). "
    "The hospital's LOCAL model is unaffected — only the on-chain submission is corrupted."
)

label("Detection Mechanism")
bullet("Byzantine weights produce poor accuracy on validator validation sets → REJECTED "
       "by >1/3 of validators → excluded from FedAvg.")
bullet("Trust score repeatedly penalized (delta_acc << -0.01 every round).")
bullet("Trust score converges to the MIN_SCORE floor (0.05) within ~5 rounds.")
bullet("Even if a Byzantine block somehow passes consensus (pBFT tolerates up to f < n/3 "
       "Byzantine validators), its low trust weight gives it <5% contribution to FedAvg.")

label("pBFT Safety Threshold")
body(
    "With 3 validators, up to f=1 Byzantine validator can be tolerated at the "
    "block level (requires 3f+1 validators for f faults → 3 validators tolerates f=0.67, "
    "so effectively no single Byzantine validator can block consensus but "
    "can influence 1/3 of vote scores)."
)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — END-TO-END DATA FLOW SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════

h1("8. End-to-End Data Flow Summary")

body("The following sequence describes one complete PoCL FL round:")

add_table(
    ["Step", "Actor",               "Action",                                          "On-Chain Record"],
    [
        ["1",  "Pneumonia Loader",   "Load 64×64 X-ray images, build train/val/test splits","—"],
        ["2",  "FL Framework",       "Initialize global MLP model (He init), create blockchain genesis block","Genesis Block"],
        ["3",  "Hospital Trainer",   "pull_global: accept/reject federated global model","—"],
        ["4",  "Hospital Trainer",   "Local SGD training (local_epochs, early stopping)","—"],
        ["5",  "Hospital Trainer",   "Sign and submit MODEL_UPDATE tx (weights hash only)","MODEL_UPDATE tx"],
        ["6",  "Hospital Trainer",   "Run inference on shared eval batch (64 samples), submit PREDICTION_PROPOSAL","PREDICTION_PROPOSAL tx"],
        ["7",  "BC Validators (x3)", "Score each hospital's predictions: accuracy + timeliness → vote_score","VOTE tx (12 total)"],
        ["8",  "FL Framework",       "Average vote scores; select top-3 winners","WINNER_SELECTION tx"],
        ["9",  "FL Framework",       "Compute R_i contribution score for each winner","REWARD tx (3 total)"],
        ["10", "FL Framework",       "Update trust scores for all hospitals","TRUST_UPDATE tx (4 total)"],
        ["11", "FL Framework",       "FedAvg(winners only) → new global model weights","—"],
        ["12", "FL Framework",       "Evaluate new global model on pooled test set","—"],
        ["13", "FL Framework",       "Commit GLOBAL_MODEL tx (hash, accuracy, trust scores)","GLOBAL_MODEL tx"],
        ["14", "PoS Consensus",      "Select proposer, run pBFT vote (>2/3 stake), seal block","Block #N committed"],
        ["15", "All Hospitals",      "Receive new global model; prepare for next round","—"],
    ],
    col_widths=[0.4, 1.7, 3.2, 1.6],
)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — EXPERIMENTAL RESULTS
# ═══════════════════════════════════════════════════════════════════════════════

h1("9. Experimental Results (50 Rounds, k=3 Winners)")

h2("9.1 Global Model Performance")
add_table(
    ["Metric",                    "Value"],
    [
        ["Final global model accuracy", "92.79%"],
        ["Chain length",                "51 blocks (1 genesis + 50 FL rounds)"],
        ["Chain validity",              "True (all blocks pass integrity check)"],
        ["Total transactions",          "1,450 (29 per round × 50 rounds)"],
        ["Tamper detection",            "True (test confirmed)"],
    ],
    col_widths=[2.8, 3.5],
)

h2("9.2 Per-Hospital Accuracy (Before vs After Federation)")
add_table(
    ["Hospital",         "Baseline (Solo)", "After FL + PoCL", "Gain",    "Objective Met?"],
    [
        ["Galle (Node A)",   "91.41%",         "88.34%",          "-3.07pp", "No"],
        ["Colombo (Node B)", "84.05%",         "93.25%",          "+9.20pp", "Yes (>5pp)"],
        ["Kandy (Node C)",   "92.64%",         "94.48%",          "+1.84pp", "No"],
        ["Jaffna (Node D)",  "92.02%",         "93.87%",          "+1.85pp", "No"],
    ],
    col_widths=[1.8, 1.5, 1.5, 1.0, 1.3],
)
note(
    "Colombo (lowest baseline hospital) benefits most from federation. "
    "High-baseline hospitals near the MLP capacity ceiling gain less because "
    "FedAvg necessarily dilutes their sharp local optima with cross-hospital averages."
)

h2("9.3 PoCL Winner Selection Statistics (50 rounds, 150 total winner slots)")
add_table(
    ["Hospital",         "Win Count", "Win Rate", "Mean R_i Contribution"],
    [
        ["Galle (Node A)",   "39",   "78.0%",  "0.0019"],
        ["Colombo (Node B)", "41",   "82.0%",  "0.0007"],
        ["Kandy (Node C)",   "41",   "82.0%",  "0.0003"],
        ["Jaffna (Node D)",  "29",   "58.0%",  "0.0020"],
    ],
    col_widths=[1.8, 1.0, 1.0, 2.0],
)
note(
    "Galle and Jaffna have higher R_i (more divergence from global model), "
    "indicating their local data distributions are most distinct from the "
    "cross-hospital average. Jaffna wins fewer rounds (58%) despite high R_i, "
    "suggesting its timeliness score is relatively lower."
)

h2("9.4 Blockchain Transaction Audit")
add_table(
    ["TX Type",             "Total Count"],
    [
        ["MODEL_UPDATE",       "200"],
        ["PREDICTION_PROPOSAL","200"],
        ["VOTE",               "600"],
        ["WINNER_SELECTION",   "50"],
        ["TRUST_UPDATE",       "200"],
        ["REWARD",             "150"],
        ["GLOBAL_MODEL",       "50"],
        ["TOTAL",              "1,450"],
    ],
    col_widths=[2.2, 1.5],
)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — FILE STRUCTURE
# ═══════════════════════════════════════════════════════════════════════════════

h1("10. Project File Structure")

add_table(
    ["File / Folder",                          "Role"],
    [
        ["core/pneumonia_loader.py",            "Data loading, preprocessing, train/val/test splitting"],
        ["core/flobc_pneumonia_engine.py",      "PneumoniaModel, HospitalTrainer, BCValidator, TrustService, FloBCPneumonia engine (train + train_pocl)"],
        ["blockchain/crypto.py",                "SHA-256, Merkle Tree, RSA-2048 Wallet"],
        ["blockchain/transaction.py",           "Transaction dataclass + 8 factory functions for all TX types"],
        ["blockchain/chain.py",                 "Block, ProofOfStake, RealBlockchain"],
        ["evaluate_objectives.py",              "Objective verification: baseline, FL, PoCL, per-hospital accuracy measurement"],
        ["run_pneumonia.py",                    "Stage 2 experiment runner: BSP/SSP/BAP sync comparison"],
        ["dashboard/plot_pneumonia.py",         "Chart generation (8 charts including PoCL winner stats)"],
        ["data/Node_A/ .. Node_D/",             "Real chest X-ray images (NORMAL + PNEUMONIA per hospital)"],
        ["results/objective_verification.json", "Full experiment results: accuracy logs, winner_log, reward_log"],
        ["results/blockchain_objectives.json",  "Full blockchain export (all 51 blocks, all transactions)"],
        ["results/objective_report.txt",        "Human-readable objective verification summary"],
        ["dashboard/pocl_winner_stats.png",     "Chart 8: PoCL winner selection frequency and R_i reward"],
    ],
    col_widths=[2.6, 4.3],
)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — BLOCK DIAGRAM GUIDE
# ═══════════════════════════════════════════════════════════════════════════════

h1("11. Block Diagram Construction Guide")

body(
    "Use the following component groups and arrows to construct the system "
    "block diagram. Recommended layout: top-to-bottom, left-to-right flow."
)

h2("Layer 1 — Data Inputs (Top)")
bullet("Box: 'Real X-Ray Images' — 4 sub-boxes for Node A/B/C/D")
bullet("Arrow down: '64x64 grayscale, normalised, flattened to 4096 features'")

h2("Layer 2 — Hospital Trainer Nodes")
bullet("4 parallel boxes: 'Hospital Trainer A', 'B', 'C', 'D'")
bullet("Inside each box: 'MLP (4096→256→2)', 'RSA Wallet', 'Local Train Data (Private)'")
bullet("Each box: arrows for pull_global IN, weights/hash OUT")

h2("Layer 3 — PoCL Consensus (6 Phases, Left to Right)")
bullet("Phase A: '4x MODEL_UPDATE tx' → Blockchain")
bullet("Phase B: '4x PREDICTION_PROPOSAL tx' → Shared Eval Batch (64 samples from Pooled Val)")
bullet("Phase C: '3 BC Validators' → '12x VOTE tx' (accuracy + timeliness)")
bullet("Phase D: 'Winner Selection' box → 'WINNER_SELECTION tx' → Top-3 selected")
bullet("Phase E: 'Reward: R_i contribution' + 'TRUST_UPDATE' → Blockchain")
bullet("Phase F: 'FedAvg (winners only)' → New Global Model")

h2("Layer 4 — Blockchain (Right side or bottom)")
bullet("'RealBlockchain' box with: 'Genesis Block', 'Block #1..#50'")
bullet("Inside each block: 'Merkle Root', 'PoS Proposer', 'stake_votes', 'block_hash'")
bullet("Show chain links: hash pointer arrows between blocks")
bullet("Show PoS consensus: '3 Validators' → '>2/3 stake = YES' → 'Block Finalized'")

h2("Layer 5 — Outputs (Bottom)")
bullet("'Global Model (updated)' → fed back to all Hospital Trainer boxes")
bullet("'objective_verification.json' → 'Dashboard Charts (8 charts)'")
bullet("'blockchain_objectives.json' → 'Chain Audit / Tamper Detection'")

h2("Key Annotations for Diagram")
add_table(
    ["Arrow / Flow",                            "Label"],
    [
        ["Hospital to blockchain",               "SHA-256(weights) only — no raw data"],
        ["Global model to hospitals",            "pull_global: accept if within 2% of local"],
        ["Validators to vote",                   "vote = 0.7×accuracy + 0.3×timeliness"],
        ["Winner selection threshold",           "Top-K = 3 of 4 per round"],
        ["FedAvg formula",                       "new_global = sum(trust_i × weights_i) for winners"],
        ["Block seal threshold",                 ">2/3 stake = pBFT consensus"],
        ["Reward formula",                       "R_i = mean|local_weights - global_weights| per layer"],
        ["Trust score bounds",                   "MIN=0.05, MAX=1.0, normalized to sum=1.0"],
    ],
    col_widths=[2.8, 4.0],
)

# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════

doc.save(OUT_PATH)
print(f"\n  Document saved → {OUT_PATH}\n")
